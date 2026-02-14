import os
from fpdf import FPDF
from docx import Document
from datetime import datetime

class ReportGenerator:
    @staticmethod
    def to_pdf(content: str, output_path: str):
        """Converts markdown-style text to a PDF file."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Add a professional header
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "HDA - Health Data Analysis Report", ln=True, align='C')
        pdf.set_font("Arial", 'I', 10)
        pdf.cell(0, 10, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True, align='R')
        pdf.ln(5)
        
        # Basic Markdown Handling (Header, Bold, List)
        # Note: This is a simplified parser for demo purposes
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(5)
                continue
                
            if line.startswith('###'):
                pdf.set_font("Arial", 'B', 14)
                pdf.multi_cell(0, 10, line.replace('###', '').strip())
            elif line.startswith('##'):
                pdf.set_font("Arial", 'B', 15)
                pdf.multi_cell(0, 10, line.replace('##', '').strip())
            elif line.startswith('#'):
                pdf.set_font("Arial", 'B', 16)
                pdf.multi_cell(0, 10, line.replace('#', '').strip())
            elif line.startswith('**') and line.endswith('**'):
                pdf.set_font("Arial", 'B', 12)
                pdf.multi_cell(0, 7, line.replace('**', '').strip())
            elif line.startswith('- ') or line.startswith('* '):
                pdf.set_font("Arial", '', 12)
                pdf.multi_cell(0, 7, f"  • {str(line)[2:]}")
            else:
                pdf.set_font("Arial", '', 12)
                # Clean up bold markers within text if any
                clean_line = str(line).replace('**', '')
                pdf.multi_cell(0, 7, clean_line)
        
        pdf.output(output_path)
        return output_path

    @staticmethod
    def to_word(content: str, output_path: str):
        """Converts markdown-style text to a Word .docx file."""
        doc = Document()
        
        # Add Header
        doc.add_heading('HDA - Health Data Analysis Report', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('###'):
                doc.add_heading(str(line).replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(str(line).replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(str(line).replace('#', '').strip(), level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(str(line)[2:], style='List Bullet')
            else:
                p = doc.add_paragraph()
                # Handle bold parts
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
        
        doc.save(output_path)
        return output_path
